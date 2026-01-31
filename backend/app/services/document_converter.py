import os
from pathlib import Path
from typing import Optional
import PyPDF2
from pdf2docx import Converter
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
import pandas as pd

class DocumentConverter:
    """Handles document format conversions"""
    
    @staticmethod
    async def convert(input_path: str, output_path: str, source_format: str, target_format: str) -> str:
        """
        Convert document from source format to target format
        
        Args:
            input_path: Path to input file
            output_path: Path to save output file
            source_format: Source file format (without dot)
            target_format: Target file format (without dot)
            
        Returns:
            Path to converted file
        """
        source_format = source_format.lower()
        target_format = target_format.lower()
        
        # PDF conversions
        if source_format == 'pdf':
            if target_format in ['doc', 'docx']:
                return await DocumentConverter.pdf_to_docx(input_path, output_path)
            elif target_format == 'txt':
                return await DocumentConverter.pdf_to_txt(input_path, output_path)
        
        # DOCX conversions
        elif source_format == 'docx':
            if target_format == 'pdf':
                return await DocumentConverter.docx_to_pdf(input_path, output_path)
            elif target_format == 'txt':
                return await DocumentConverter.docx_to_txt(input_path, output_path)
        
        # TXT conversions
        elif source_format == 'txt':
            if target_format == 'pdf':
                return await DocumentConverter.txt_to_pdf(input_path, output_path)
            elif target_format == 'docx':
                return await DocumentConverter.txt_to_docx(input_path, output_path)
        
        # Excel conversions
        elif source_format in ['xls', 'xlsx']:
            if target_format == 'csv':
                return await DocumentConverter.excel_to_csv(input_path, output_path)
            elif target_format == 'pdf':
                return await DocumentConverter.excel_to_pdf(input_path, output_path)
        
        # CSV conversions
        elif source_format == 'csv':
            if target_format == 'xlsx':
                return await DocumentConverter.csv_to_excel(input_path, output_path)
            elif target_format == 'pdf':
                return await DocumentConverter.csv_to_pdf(input_path, output_path)
        
        raise NotImplementedError(f"Conversion from {source_format} to {target_format} is not yet implemented")
    
    @staticmethod
    async def pdf_to_docx(input_path: str, output_path: str) -> str:
        """Convert PDF to DOCX"""
        cv = Converter(input_path)
        cv.convert(output_path)
        cv.close()
        return output_path
    
    @staticmethod
    async def pdf_to_txt(input_path: str, output_path: str) -> str:
        """Convert PDF to TXT"""
        text_content = []
        
        with open(input_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
        
        with open(output_path, 'w', encoding='utf-8') as output_file:
            output_file.write('\n\n'.join(text_content))
        
        return output_path
    
    @staticmethod
    async def docx_to_pdf(input_path: str, output_path: str) -> str:
        """Convert DOCX to PDF using reportlab"""
        # Read DOCX
        doc = Document(input_path)
        
        # Create PDF
        pdf = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Extract text from DOCX and add to PDF
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                p = Paragraph(paragraph.text, styles['Normal'])
                story.append(p)
                story.append(Spacer(1, 0.2 * inch))
        
        pdf.build(story)
        return output_path
    
    @staticmethod
    async def docx_to_txt(input_path: str, output_path: str) -> str:
        """Convert DOCX to TXT"""
        doc = Document(input_path)
        
        with open(output_path, 'w', encoding='utf-8') as output_file:
            for paragraph in doc.paragraphs:
                output_file.write(paragraph.text + '\n')
        
        return output_path
    
    @staticmethod
    async def txt_to_pdf(input_path: str, output_path: str) -> str:
        """Convert TXT to PDF"""
        with open(input_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        pdf = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                # Replace line breaks within paragraphs
                para_text = para_text.replace('\n', ' ')
                p = Paragraph(para_text, styles['Normal'])
                story.append(p)
                story.append(Spacer(1, 0.2 * inch))
        
        pdf.build(story)
        return output_path
    
    @staticmethod
    async def txt_to_docx(input_path: str, output_path: str) -> str:
        """Convert TXT to DOCX"""
        doc = Document()
        
        with open(input_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text)
        
        doc.save(output_path)
        return output_path
    
    @staticmethod
    async def excel_to_csv(input_path: str, output_path: str) -> str:
        """Convert Excel to CSV"""
        # Read first sheet
        df = pd.read_excel(input_path, sheet_name=0)
        df.to_csv(output_path, index=False)
        return output_path
    
    @staticmethod
    async def csv_to_excel(input_path: str, output_path: str) -> str:
        """Convert CSV to Excel"""
        df = pd.read_csv(input_path)
        df.to_excel(output_path, index=False, engine='openpyxl')
        return output_path
    
    @staticmethod
    async def excel_to_pdf(input_path: str, output_path: str) -> str:
        """Convert Excel to PDF"""
        # Read Excel
        df = pd.read_excel(input_path, sheet_name=0)
        
        # Create PDF
        pdf = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add table data as text (simplified)
        for idx, row in df.iterrows():
            row_text = ' | '.join([str(val) for val in row])
            p = Paragraph(row_text, styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 0.1 * inch))
        
        pdf.build(story)
        return output_path
    
    @staticmethod
    async def csv_to_pdf(input_path: str, output_path: str) -> str:
        """Convert CSV to PDF"""
        # Read CSV
        df = pd.read_csv(input_path)
        
        # Create PDF
        pdf = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add headers
        headers = ' | '.join(df.columns)
        story.append(Paragraph(headers, styles['Heading2']))
        story.append(Spacer(1, 0.2 * inch))
        
        # Add data rows
        for idx, row in df.iterrows():
            row_text = ' | '.join([str(val) for val in row])
            p = Paragraph(row_text, styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 0.1 * inch))
        
        pdf.build(story)
        return output_path
